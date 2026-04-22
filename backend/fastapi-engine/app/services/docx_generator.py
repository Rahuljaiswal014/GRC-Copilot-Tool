"""
DOCX report generator using python-docx.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger


def generate_docx(report_data: dict, output_dir: str = "./reports") -> str:
    """Generate a DOCX report and return the file path."""
    os.makedirs(output_dir, exist_ok=True)

    filename = f"grc_report_{report_data.get('assessment_id', 'unknown')}.docx"
    filepath = os.path.join(output_dir, filename)

    doc = Document()

    # Title
    title = doc.add_heading('GRC Assessment Report', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta
    meta = doc.add_paragraph()
    meta.add_run(
        f"Framework: {report_data.get('framework', 'N/A')} | "
        f"Organization: {report_data.get('organization', 'N/A')} | "
        f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}"
    ).font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Executive Summary
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(report_data.get('executive_summary', ''))

    score = report_data.get('compliance_score', 0)
    risk = report_data.get('risk_level', 'N/A')
    p = doc.add_paragraph()
    p.add_run(f"Compliance Score: {score}%").bold = True
    p.add_run(f" | Risk Level: {risk}")

    # Domain Scores
    doc.add_heading('Domain Scores', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'

    headers = ['Domain', 'Score', 'Total', 'Answered']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header

    for d in report_data.get('domain_scores', []):
        row = table.add_row()
        row.cells[0].text = d.get('domain', '')
        row.cells[1].text = f"{d.get('score', 0)}%"
        row.cells[2].text = str(d.get('questions_total', 0))
        row.cells[3].text = str(d.get('questions_answered', 0))

    doc.add_paragraph()

    # Recommendations
    doc.add_heading('Recommendations', level=2)
    for rec in report_data.get('recommendations', []):
        p = doc.add_paragraph()
        p.add_run(f"{rec.get('title', '')} ").bold = True
        p.add_run(f"[{rec.get('priority', '').upper()}] ({rec.get('horizon', '')})")
        doc.add_paragraph(rec.get('detail', ''))
        doc.add_paragraph(
            f"Est. Cost: INR {rec.get('cost_inr', 0):,} | Effort: {rec.get('effort', '')}"
        )

    # Cost Summary
    doc.add_heading('Cost Summary', level=2)
    cost = report_data.get('cost_summary', {})
    doc.add_paragraph(f"Total Estimated Cost: INR {cost.get('total_inr', 0):,}")
    doc.add_paragraph(f"Critical Items: INR {cost.get('critical_inr', 0):,}")

    doc.save(filepath)
    logger.info(f"DOCX generated: {filepath}")
    return filepath
